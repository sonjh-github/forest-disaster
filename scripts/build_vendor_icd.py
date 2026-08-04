from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "vendor-icd"
QA = OUT / "qa"
OUT.mkdir(parents=True, exist_ok=True)
QA.mkdir(parents=True, exist_ok=True)

TODAY = "2026년 8월 3일"
FONT = "맑은 고딕"
MONO = "Consolas"
NAVY = "1F4E78"
BLUE = "2E75B6"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE5"
TEXT_GRAY = "555555"
RISK = "9B1C1C"
WHITE = "FFFFFF"
PAGE_BREAK_PENDING = False


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=75, start=120, bottom=75, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_font(run, name=FONT, size=10.5, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_para(p, before=0, after=5, line=1.18, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep


def add_text(doc, text, *, size=10.5, bold=False, color=None, align=None, before=0, after=5, italic=False):
    p = doc.add_paragraph()
    set_para(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    global PAGE_BREAK_PENDING
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    if PAGE_BREAK_PENDING:
        p.paragraph_format.page_break_before = True
        PAGE_BREAK_PENDING = False
    r = p.add_run(text)
    set_font(r, size={1: 15.5, 2: 12.5, 3: 11.5}[level], bold=True, color=NAVY if level == 1 else BLUE)
    return p


def add_callout(doc, title: str, body: str, risk=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FCE8E6" if risk else LIGHT_BLUE)
    p = cell.paragraphs[0]
    set_para(p, after=2)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=RISK if risk else NAVY)
    p2 = cell.add_paragraph()
    set_para(p2, after=0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.2)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=0)
        r = p.add_run(text)
        set_font(r, size=9.2, bold=True, color=WHITE)
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.12)
            if idx == 0 and len(headers) <= 3:
                r = p.add_run(str(text))
                set_font(r, size=font_size, bold=True, color=NAVY)
            else:
                r = p.add_run(str(text))
                set_font(r, size=font_size)
    return table


def add_code(doc, payload: dict):
    text = json.dumps(payload, ensure_ascii=False, indent=2)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.0)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, name=MONO, size=8.1, color="1F2937")


def add_page_break(doc):
    global PAGE_BREAK_PENDING
    PAGE_BREAK_PENDING = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=TEXT_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_doc(recipient: str, title: str, scope: str):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for level, size, before, after in ((1, 15.5, 15, 8), (2, 12.5, 11, 6), (3, 11.5, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(hp, after=0)
    hr = hp.add_run("투비유니콘 | 산불 통합관제 상위 시스템 연계(Northbound) 협의안")
    set_font(hr, size=8.5, bold=True, color=TEXT_GRAY)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    add_text(doc, "투비유니콘", size=11, bold=True, color=BLUE, after=10)
    add_text(doc, title, size=23, bold=True, color=NAVY, after=5)
    add_text(doc, scope, size=12.5, color=TEXT_GRAY, after=16)
    add_table(doc, ["구분", "내용"], [
        ["수신기관", recipient],
        ["발신기관", "주식회사 투비유니콘"],
        ["문서버전", "v0.9 (업체 검토용)"],
        ["작성일", TODAY],
        ["적용구간", "업체 게이트웨이·관리 SW·네트워크 제어기 ↔ 투비유니콘 통합 API"],
        ["적용 제외", "현장 RF 구간 및 장비 내부 southbound 프로토콜"],
    ], [1900, 7460], font_size=9.6)
    add_callout(
        doc,
        "중요 - 본 문서의 성격",
        "본 문서는 장비 자체가 HTTP/JSON을 직접 처리하도록 강제하는 규격이 아닙니다. "
        "현장 장비의 원시 데이터를 수집하는 업체 측 게이트웨이·관리 SW·제어기가 투비유니콘 통합 API와 연동하기 위한 Northbound 기준안입니다. "
        "지원이 어려운 항목은 대체 인터페이스와 함께 회신해 주십시오.",
    )
    add_heading(doc, "검토 요청", 1)
    add_table(doc, ["확인 항목", "업체 회신"], [
        ["HTTPS/JSON Push 구현 가능 여부", "□ 가능  □ 조건부 가능  □ 불가"],
        ["실제 송수신 주체", "장비명·프로그램명: __________________________"],
        ["제안안 수정 필요 여부", "□ 없음  □ 있음 — 수정안 또는 업체 규격 별첨"],
        ["대체 상위 연계방식", "□ 업체 REST API  □ MQTT  □ TCP Socket  □ SNMP/Trap  □ 기타 IP 방식"],
        ["기술 담당자", "성명/연락처: _________________________________"],
    ], [4100, 5260], font_size=9.5)
    add_text(doc, "※ 검토 완료 후 상호 협의된 내용을 반영하여 v1.0 확정본을 발행합니다.", size=9, color=TEXT_GRAY, italic=True)
    return doc


def add_common_sections(doc, vendor_system: str, reporting_role: str):
    add_page_break(doc)
    add_heading(doc, "1. 연동 원칙과 책임 경계", 1)
    add_table(doc, ["구간", "담당", "원칙"], [
        ["현장 장비 ↔ 업체 제어기", "수신 업체", "업체 고유 프로토콜 사용. 투비유니콘은 원시 RF·Serial 프레임을 규정하지 않음"],
        ["업체 제어기 ↔ 통합 API", "공동", "HTTPS POST 및 application/json을 기준으로 협의"],
        ["JSON 변환·전송·재전송", "수신 업체", "업체 측 게이트웨이·관리 SW·제어기에서 수행"],
        ["API 수신·검증·저장·표출", "투비유니콘", "UUID 검증, 중복 방지, DB 저장, 상황판 표출"],
        ["자산 UUID 발급·매핑", "공동", "투비유니콘 발급 UUID와 업체 고유번호 매핑표 관리"],
    ], [2200, 1600, 5560], font_size=9.2)
    add_callout(doc, "권장 전송 구조", f"현장 장비 → {vendor_system} → HTTPS/JSON → 투비유니콘 통합 API")

    add_heading(doc, "2. 접속 기준", 1)
    add_table(doc, ["항목", "기준"], [
        ["운영 Base URL", "상호 협의 후 별도 제공 (예: https://{host})"],
        ["전송 방식", "업체 시스템에서 통합 API로 HTTPS POST Push"],
        ["Content-Type", "application/json; charset=utf-8"],
        ["인증", "사전 등록된 게이트웨이 자격증명으로 Bearer Token 발급"],
        ["중복 방지", "Idempotency-Key 헤더 = context.requestId"],
        ["시각", "ISO 8601. 운영 권장값은 UTC(Z), KST 오프셋(+09:00)도 허용 협의"],
        ["좌표", "GeoJSON Point, 순서 [경도, 위도, 고도(m)]"],
        ["보안", "운영환경은 HTTPS 필수. mTLS·기관 PKI는 보안협의 후 확정"],
    ], [2400, 6960], font_size=9.4)

    add_heading(doc, "3. 인증 및 공통 요청 형식", 1)
    add_text(doc, "3.1 게이트웨이 활성화", size=11.5, bold=True, color=BLUE, before=3, after=4)
    add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/gateways/activate"],
        ["목적", "15분 유효 Bearer Token 발급"],
    ], [2000, 7360], font_size=9.4)
    add_code(doc, {
        "assetId": "20000000-0000-4000-8000-000000000001",
        "credential": "<별도 발급 자격증명>",
        "eventId": "10000000-0000-4000-8000-000000000001",
        "reportingRole": reporting_role,
    })
    add_text(doc, "3.2 결과 전송 공통 헤더", size=11.5, bold=True, color=BLUE, before=5, after=4)
    add_table(doc, ["헤더", "값", "필수"], [
        ["Authorization", "Bearer <accessToken>", "예"],
        ["Content-Type", "application/json", "예"],
        ["Idempotency-Key", "context.requestId와 동일한 UUID", "예"],
        ["X-Origin", "업체 시스템 식별자", "권장"],
    ], [2500, 5200, 1660], font_size=9.1)

    add_heading(doc, "4. 공통 Envelope", 1)
    add_table(doc, ["필드", "형식", "설명"], [
        ["context.eventId", "UUID", "운영·실증 사건 식별자"],
        ["context.requestId", "UUID", "메시지 고유번호. 재전송 시 동일 값 유지"],
        ["context.sourceSystem", "문자열", "송신 프로그램·시스템 식별자"],
        ["context.occurredAt", "date-time", "데이터 발생시각"],
        ["context.sentAt", "date-time", "실제 전송시각(권장)"],
        ["context.schemaVersion", "문자열", "본 문서는 1.0"],
        ["context.reportedByAssetId", "UUID", "실제 HTTP 호출 주체 UUID"],
        ["context.reportingRole", "열거형", f"본 문서 권장값: {reporting_role}"],
        ["data", "객체", "장비유형별 결과 데이터"],
    ], [3300, 1800, 4260], font_size=8.9)


def add_delivery_rules(doc):
    add_heading(doc, "전송·응답·재전송 규칙", 1)
    add_table(doc, ["상황", "처리 기준"], [
        ["정상 수신", "HTTP 200 및 data.accepted=true 확인"],
        ["동일 요청 재전송", "동일 Idempotency-Key와 requestId 유지. 서버는 중복 저장하지 않음"],
        ["HTTP 400", "필드·형식 오류 수정 후 재전송. 동일 데이터면 requestId 유지"],
        ["HTTP 401/403", "토큰 재발급, 자산·사건 배정 및 보고 권한 확인"],
        ["HTTP 5xx/Timeout", "로컬 버퍼 보관 후 지수 백오프 재시도"],
        ["현장망 단절", "발생시각·원 sequence를 보존하여 복구 후 순차 전송"],
        ["전송주기", "업체 지원주기 확인 후 v1.0에서 확정. 상태변경 이벤트는 즉시 전송 권장"],
    ], [2500, 6860], font_size=9.2)

    add_heading(doc, "제안한 HTTPS/JSON 적용이 어려운 경우", 1)
    add_callout(
        doc,
        "대체 상위연계 검토",
        "게이트웨이·관리 SW·네트워크 제어기가 제안한 HTTPS/JSON Push를 적용하기 어려운 경우, "
        "업체 REST API, MQTT, TCP Socket 등 실제 제공 가능한 IP 기반 상위 연계방식과 수정안을 회신해 주십시오. "
        "Serial·Modbus 등 장비 내부 또는 현장 하위 인터페이스는 통합시스템의 Northbound 대안으로 보지 않습니다. "
        "필요한 변환 기능은 업체 게이트웨이·관리 SW 또는 별도 연동 어댑터의 책임 범위로 협의합니다.",
        risk=True,
    )
    add_table(doc, ["대체 방식", "제공 요청사항"], [
        ["MQTT", "Broker/TLS, Topic, QoS, Payload 샘플"],
        ["TCP Socket", "접속 방향, 프레임 구조, IP/Port, ACK·재전송 규칙"],
        ["업체 REST API", "OpenAPI/Swagger, 인증, 조회주기, 오류코드"],
        ["SNMP/Trap", "게이트웨이·망 관리상태에 한해 MIB, OID, Trap 명세"],
    ], [2100, 7260], font_size=9.2)


def add_vendor_decision_pages(doc, vendor: str, transmitter_options: str, review_focus: str):
    """Request one Northbound method and a concrete vendor revision proposal."""
    add_page_break(doc)
    add_heading(doc, "업체 회신 요청 — 연동방식 및 수정안", 1)
    add_callout(
        doc,
        "회신 기준",
        "현장 단말의 원시 RF 또는 장비 내부 인터페이스가 아니라, "
        "업체가 납품하는 게이트웨이·관리 SW·네트워크 제어기에서 투비유니콘 통합시스템으로 제공할 수 있는 "
        "상위 시스템 연계(Northbound) 인터페이스를 기준으로 검토해 주십시오. "
        "본 문서의 HTTPS/JSON은 투비유니콘 제안안이며, 적용이 어려운 경우 업체 수정안을 회신해 주십시오.",
    )
    add_callout(doc, "업체별 검토 대상", review_focus)

    add_heading(doc, "1. 주 연동방식 검토", 2)
    add_table(doc, ["선택", "방식", "회신 요청사항"], [
        ["□", "A. 제안안대로 HTTPS/JSON Push", "제시된 접속·인증·Payload·응답 규칙 사용"],
        ["□", "B. 업체 수정 HTTPS/JSON Push", "변경이 필요한 항목과 수정 JSON 샘플 제시"],
        ["□", "C. 업체 REST API 제공", "투비유니콘이 조회할 API 명세와 응답 샘플 제시"],
        ["□", "D. MQTT Publish", "Broker, TLS, Topic, QoS와 Payload 샘플 제시"],
        ["□", "E. TCP Socket", "접속 방향, IP/Port, Frame, ACK와 재전송 규칙 제시"],
        ["□", "F. SNMP/Trap", "게이트웨이·망 관리상태에 한해 MIB·OID·Trap 명세 제시"],
        ["□", "G. 다른 IP 기반 방식", "방식명과 인터페이스 명세·데이터 샘플 제시"],
        ["□", "H. 현재 구조로 연계 곤란", "기술적 사유와 가능한 구조 변경안을 제시"],
    ], [650, 3050, 5660], font_size=8.6)

    add_heading(doc, "2. 실제 송신·제공 주체 확정", 2)
    add_table(doc, ["항목", "업체 확정 내용"], [
        ["후보 주체", transmitter_options],
        ["최종 장비·프로그램명", ""],
        ["제품/펌웨어/SW 버전", ""],
        ["설치 위치", "□ 현장 장비  □ 지휘차량  □ 업체 서버  □ 통합서버  □ 기타: __________"],
        ["접속 방향", "□ 업체→투비 Push  □ 투비→업체 Pull  □ 양방향"],
        ["외부망 접속", "□ 가능  □ 방화벽·VPN 협의 필요  □ 불가"],
    ], [2700, 6660], font_size=9.0)

    add_page_break(doc)
    add_heading(doc, "2. 업체 수정안 작성", 1)
    add_text(doc, "제안안과 다른 부분만 작성해도 됩니다. 수정 범위가 큰 경우 업체 문서 또는 별첨 규격으로 회신해 주십시오.", after=7)
    add_table(doc, ["검토 항목", "투비유니콘 제안", "업체 수정안·의견"], [
        ["실제 연계 주체", "게이트웨이·관리 SW·네트워크 제어기", ""],
        ["연결 방향", "업체 시스템 → 통합 API Push", ""],
        ["전송 방식", "HTTPS POST / application/json", ""],
        ["접속 주소", "투비유니콘이 운영 URL 제공", ""],
        ["인증·보안", "Bearer Token, 운영환경 HTTPS", ""],
        ["메시지 구조", "공통 Envelope + 장비·망 데이터", ""],
        ["이벤트·전송주기", "상태변경 즉시, 주기보고는 상호 협의", ""],
        ["응답·중복방지", "HTTP 응답, requestId·Idempotency-Key", ""],
        ["재전송·망 단절", "로컬 버퍼 후 원 시각을 유지하여 재전송", ""],
        ["식별자", "투비유니콘 UUID와 업체 장비 ID 매핑", ""],
        ["시간·좌표·단위", "ISO 8601, GeoJSON, SI 단위", ""],
        ["제어 명령", "지원 명령 확인 후 별도 규격으로 협의", ""],
    ], [2450, 3500, 3410], font_size=8.5)

    add_heading(doc, "3. 업체 수정 Payload 또는 응답 샘플", 2)
    add_text(doc, "JSON 또는 사용 예정 형식의 샘플을 아래에 붙이거나 별첨 파일로 제출해 주십시오.", after=4)
    add_table(doc, ["구분", "업체 작성"], [
        ["샘플 데이터 또는 별첨 파일명", ""],
        ["필수 전제조건·제약사항", ""],
        ["추가 협의가 필요한 사항", ""],
    ], [2900, 6460], font_size=9.0)

    add_heading(doc, "4. 함께 요청하는 자료", 2)
    add_table(doc, ["요청자료", "회신"], [
        ["실제 송수신 Payload·응답 또는 Packet 샘플", "□ 본문 기재  □ 별첨: ______________________"],
        ["API·MQTT·TCP·SNMP 중 선택한 방식의 명세", "□ 별첨: _________________________________"],
        ["오류응답·ACK·Timeout·재전송 규칙", "□ 본문 기재  □ 별첨: ______________________"],
        ["송수신 주체와 연결 방향이 표시된 구성도", "□ 별첨: _________________________________"],
        ["기술 담당자", "성명/연락처/이메일: __________________________"],
    ], [4800, 4560], font_size=8.8)


def add_review_page(doc, vendor: str, questions: list[str]):
    add_page_break(doc)
    add_heading(doc, "세부 기능 검토 체크리스트", 1)
    add_text(doc, f"수신기관: {vendor}", bold=True, after=3)
    add_text(doc, "연동방식 확정서의 결론과 일치하도록 각 항목을 모두 표시해 주십시오.", after=8)
    rows = [[str(i + 1), q, "□ 가능\n□ 조건부\n□ 불가", ""] for i, q in enumerate(questions)]
    add_table(doc, ["No.", "확인사항", "판정", "의견·조건"], rows, [650, 4860, 1450, 2400], font_size=8.8)
    add_heading(doc, "송수신 주체 확인", 2)
    add_table(doc, ["항목", "회신 내용"], [
        ["장비 또는 프로그램명", ""],
        ["운영체제·펌웨어", ""],
        ["설치 위치", "□ 지휘차량  □ 현장 장비  □ 업체 서버  □ 기타"],
        ["외부 HTTPS 접속", "□ 가능  □ 방화벽 협의 필요  □ 불가"],
        ["예상 전송주기", ""],
        ["기술 담당자", ""],
    ], [2900, 6460], font_size=9.3)
    add_page_break(doc)
    add_heading(doc, "검토 결과", 2)
    add_table(doc, ["구분", "기재"], [
        ["종합 의견", "□ 현 규격 구현 가능   □ 수정 후 가능   □ 대체 인터페이스 협의 필요"],
        ["주요 수정요청", ""],
        ["회신일", "년       월       일"],
        ["검토자", "소속/직위:                         성명:                    (서명)"],
    ], [2500, 6860], font_size=9.4)
    add_text(doc, "※ 본 회신은 기술 협의를 위한 것으로, 최종 규격은 상호 합의된 v1.0 문서로 확정합니다.", size=9, color=TEXT_GRAY, italic=True, before=5)


def build_ndps():
    doc = setup_doc(
        "앤디피에스 주식회사",
        "TVWS 상위 시스템 연계(Northbound) 인터페이스 협의안",
        "TVWS Base/CPE 및 이동기지국 관리 SW 연계 | HTTPS/JSON 제안 | v0.9 검토용",
    )
    add_common_sections(doc, "TVWS Base/CPE → 이동기지국 통신모듈·관리 SW", "NMS")

    add_page_break(doc)
    add_heading(doc, "5. 엔디피에스 적용 범위", 1)
    add_table(doc, ["포함", "제외 또는 별도 협의"], [
        ["TVWS Base-CPE 링크 상태", "TVWS RF·MAC 내부 프레임"],
        ["채널, 신호세기, 처리량, 지연, 손실률", "400MHz 무전기 내부 프로토콜"],
        ["Ethernet 인입 및 외부 백홀 상태", "LPWA·RTK 단말 데이터"],
        ["연결 단말 수와 장애상태", "AI-RAN 세부 알고리즘 결과"],
        ["관리 SW·통신모듈의 상태 전송", "장비 제어 명령은 명령목록 수령 후 별도 확정"],
    ], [4680, 4680], font_size=9.3)
    add_callout(doc, "송신 주체", "TVWS Base/CPE가 직접 HTTP를 호출할 필요는 없습니다. 이동기지국 내부 통신모듈, 관리 SW 또는 별도 NMS 중 실제 구현 가능한 주체를 엔디피에스가 지정합니다.")

    add_heading(doc, "6. TVWS 링크 결과 API", 1)
    add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/integrations/wildfire.tvws-network/results"],
        ["Direction", "엔디피에스 시스템 → 투비유니콘 통합 API"],
        ["권장 이벤트", "주기 상태보고, 링크 연결·열화·단절·복구, 채널 변경"],
    ], [2300, 7060], font_size=9.4)
    fields = [
        ["baseAssetId", "UUID", "필수", "TVWS Base 통합 자산 UUID"],
        ["cpeAssetId", "UUID", "필수", "TVWS CPE 통합 자산 UUID"],
        ["observedAt", "date-time", "필수", "장비 상태 관측시각"],
        ["operationalStatus", "enum", "필수", "ONLINE / DEGRADED / OFFLINE"],
        ["channel", "string", "제안 필수", "운용 채널. 실제 표현방식 회신 필요"],
        ["signalStrengthDbm", "number", "제안 필수", "수신신호 세기(dBm)"],
        ["throughputMbps", "number", "제안 필수", "측정 처리량(Mbps)"],
        ["latencyMs", "number", "제안 필수", "링크 지연(ms)"],
        ["packetLossPct", "number", "제안 필수", "패킷 손실률(0~100%)"],
        ["distanceM", "number", "협의", "Base-CPE 거리(m). 미지원 시 산출주체 협의"],
        ["ingressMedium", "enum", "필수", "현재 기준 ETHERNET"],
        ["backhaulType", "enum", "필수", "TVWS / LTE / 5G / LEO / ETHERNET"],
        ["backhaulAvailable", "boolean", "필수", "외부 백홀 사용 가능 여부"],
        ["connectedTerminals", "integer", "제안 필수", "연결 단말 수(0 이상)"],
        ["attributes", "object", "선택", "제조사 오류코드·온도·전원 등 확장정보"],
    ]
    add_table(doc, ["필드", "형식", "구분", "설명"], fields, [2450, 1300, 1350, 4260], font_size=8.6)

    add_page_break(doc)
    add_heading(doc, "7. TVWS 결과 전송 예시", 1)
    add_code(doc, {
        "context": {
            "eventId": "10000000-0000-4000-8000-000000000001",
            "requestId": "9aa4bd80-9baf-4abd-a8da-0e63a14585af",
            "sourceSystem": "ndps-tvws-manager",
            "occurredAt": "2026-08-03T04:30:00Z",
            "sentAt": "2026-08-03T04:30:01Z",
            "schemaVersion": "1.0",
            "reportedByAssetId": "20000000-0000-4000-8000-000000000002",
            "reportingRole": "NMS"
        },
        "data": {
            "baseAssetId": "20000000-0000-4000-8000-000000000011",
            "cpeAssetId": "20000000-0000-4000-8000-000000000012",
            "observedAt": "2026-08-03T04:30:00Z",
            "operationalStatus": "ONLINE",
            "channel": "27",
            "signalStrengthDbm": -66,
            "throughputMbps": 31.4,
            "latencyMs": 48,
            "packetLossPct": 0.4,
            "distanceM": 1800,
            "ingressMedium": "ETHERNET",
            "backhaulType": "5G",
            "backhaulAvailable": True,
            "connectedTerminals": 7,
            "attributes": {"vendorStatusCode": "NORMAL"}
        }
    })
    add_callout(doc, "제어 명령 처리", "채널 변경, 출력 조절, 재기동 등의 명령 API는 엔디피에스가 실제 지원 명령·파라미터·ACK·오류코드를 회신한 뒤 v1.0 부록으로 확정합니다. 본 v0.9에서는 상태 수신 API만 우선 검토합니다.")
    add_delivery_rules(doc)
    add_vendor_decision_pages(
        doc,
        "앤디피에스 주식회사",
        "이동기지국 통신모듈 / 관리 SW / NMS / 별도 연동 프로그램 중 선택",
        "TVWS Base/CPE 링크상태, 채널, 신호세기, 처리량, 지연, 손실률, 외부 백홀, "
        "연결 단말 수와 장애·복구 이벤트를 상위 시스템에 제공하는 방식을 검토합니다.",
    )
    path = OUT / "투비유니콘_앤디피에스_TVWS_상위연계_Northbound_인터페이스_협의안_v0.9.docx"
    doc.save(path)
    return path


def build_jininfra():
    doc = setup_doc(
        "주식회사 진인프라",
        "RTK·LPWA 상위 시스템 연계(Northbound) 인터페이스 협의안",
        "RTK 단말, 이동형 기준국, LPWA RF 게이트웨이·네트워크 제어기 연계 | HTTPS/JSON 제안 | v0.9 검토용",
    )
    add_common_sections(doc, "RTK 단말 → LPWA RF 게이트웨이·네트워크 제어기", "GATEWAY")

    add_page_break(doc)
    add_heading(doc, "5. 진인프라 적용 범위", 1)
    add_table(doc, ["포함", "제외 또는 별도 협의"], [
        ["대원용 RTK 단말의 위치·측위 품질", "LPWA RF 원시 프레임 자체"],
        ["LPWA 기본망 및 LTE 보조망 전환상태", "LoRa PHY/MAC 내부 구현"],
        ["이동형 RTK 기준국 RTCM 제공상태", "RTCM 원문 전체를 JSON에 포함하는 방식"],
        ["LPWA 게이트웨이 채널·슬롯·연결 단말", "FOTA 바이너리 전달 규격"],
        ["Ethernet·LTE 등 외부 백홀 상태", "제어 명령은 명령목록 수령 후 별도 확정"],
    ], [4680, 4680], font_size=9.3)
    add_callout(doc, "송신 주체", "RTK 단말이 LPWA 구간에서 JSON을 직접 보낼 필요는 없습니다. LPWA RF 게이트웨이의 네트워크 제어기 또는 연동 인터페이스 SW가 단말 패킷을 수집·정규화하여 HTTPS/JSON으로 전송합니다.")

    add_heading(doc, "6. RTK 단말 위치 결과 API", 1)
    add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/integrations/wildfire.rtk-terminal/results"],
        ["Direction", "진인프라 게이트웨이·제어기 → 투비유니콘 통합 API"],
        ["권장 이벤트", "주기 위치보고, RTK FIX 변경, LPWA/LTE 전환, 비상상태"],
    ], [2300, 7060], font_size=9.4)
    terminal_fields = [
        ["personExternalId", "string", "필수", "대원 외부 식별자. 배정정보와 일치"],
        ["sourceAssetId", "UUID", "필수", "RTK 단말 통합 자산 UUID"],
        ["observedAt", "date-time", "필수", "측위시각"],
        ["transmittedAt", "date-time", "필수", "단말·게이트웨이 전송시각"],
        ["geometry", "GeoJSON", "필수", "Point [경도, 위도, 고도(m)]"],
        ["positioningMethod", "enum", "필수", "RTK_FIXED / RTK_FLOAT / GNSS"],
        ["horizontalAccuracyM", "number", "제안 필수", "수평 정확도(m)"],
        ["gnssFixQuality", "string", "권장", "제조사 GNSS fix 품질"],
        ["primaryLink", "enum", "필수", "LPWA"],
        ["fallbackLink", "enum", "필수", "LTE"],
        ["activeLink", "enum", "필수", "LPWA / LTE"],
        ["fallbackActivated", "boolean", "필수", "activeLink=LTE일 때 true"],
        ["signalStrengthDbm", "number", "권장", "현재 활성망 수신신호 세기"],
        ["batteryPercent", "number", "권장", "0~100"],
        ["emergency", "boolean", "권장", "비상버튼·긴급상태"],
        ["sourceSystem", "string", "필수", "원천 시스템 식별자"],
    ]
    add_table(doc, ["필드", "형식", "구분", "설명"], terminal_fields, [2450, 1300, 1350, 4260], font_size=8.5)
    add_callout(doc, "RTK 신뢰도 표시", "positioningMethod가 RTK_FIXED가 아니거나 기준국의 rtcmAvailable=false인 경우 통합관제에서는 '보정 필요/정밀 위치로 신뢰할 수 없음' 경고를 표시합니다. 보정 적용 여부를 더 명확히 제공할 수 있다면 지원 필드를 회신해 주십시오.")

    add_heading(doc, "7. RTK 단말 결과 전송 예시", 1)
    add_code(doc, {
        "context": {
            "eventId": "10000000-0000-4000-8000-000000000001",
            "requestId": "7d401a25-3b0e-4a48-9c77-dfd27a22b774",
            "sourceSystem": "jininfra-lpwa-controller",
            "occurredAt": "2026-08-03T04:30:00Z",
            "sentAt": "2026-08-03T04:30:01Z",
            "schemaVersion": "1.0",
            "reportedByAssetId": "30000000-0000-4000-8000-000000000001",
            "reportingRole": "GATEWAY"
        },
        "data": {
            "personExternalId": "WF-LEADER-001",
            "sourceAssetId": "30000000-0000-4000-8000-000000000021",
            "observedAt": "2026-08-03T04:30:00Z",
            "transmittedAt": "2026-08-03T04:30:01Z",
            "geometry": {"type": "Point", "coordinates": [128.6901, 36.3512, 312.4]},
            "positioningMethod": "RTK_FIXED",
            "horizontalAccuracyM": 0.04,
            "gnssFixQuality": "FIXED",
            "primaryLink": "LPWA",
            "fallbackLink": "LTE",
            "activeLink": "LPWA",
            "fallbackActivated": False,
            "signalStrengthDbm": -91,
            "batteryPercent": 82,
            "emergency": False,
            "sourceSystem": "rtk-terminal"
        }
    })

    add_heading(doc, "8. RTK 기준국·LPWA 게이트웨이 상태 API", 1)
    add_table(doc, ["항목", "값"], [
        ["Method", "POST"],
        ["Path", "/api/v1/integrations/wildfire.rtk-base-lpwa-gateway/results"],
        ["Direction", "진인프라 게이트웨이·제어기 → 투비유니콘 통합 API"],
        ["권장 이벤트", "기준국 가동·열화·정지, RTCM 유무, 단말 접속·백홀 변경"],
    ], [2300, 7060], font_size=9.4)
    gateway_fields = [
        ["assetId", "UUID", "필수", "기준국·LPWA 게이트웨이 통합 UUID"],
        ["observedAt", "date-time", "필수", "상태 관측시각"],
        ["operationalStatus", "enum", "필수", "ONLINE / DEGRADED / OFFLINE"],
        ["basePosition", "GeoJSON", "권장", "기준국 Point [경도, 위도, 고도]"],
        ["rtcmFormat", "string", "필수", "예: RTCM3. 실제 버전·메시지 형식 회신"],
        ["rtcmAvailable", "boolean", "필수", "보정정보 제공 가능 여부"],
        ["correctionAgeSeconds", "number", "제안 필수", "보정정보 경과시간(초)"],
        ["deliveryMode", "enum", "필수", "BROADCAST / MULTICAST"],
        ["beaconChannel", "integer", "필수", "비콘 채널"],
        ["uplinkChannelCount", "integer", "필수", "상향 채널 수"],
        ["connectedTerminals", "integer", "필수", "접속 단말 수"],
        ["allocatedSlots", "array", "필수", "단말 UUID·채널·슬롯 배정 목록"],
        ["ethernetBackhaul", "object", "필수", "연결 여부, 종류, 지연 등"],
        ["attributes", "object", "선택", "업체 오류코드·전원·온도 등"],
    ]
    add_table(doc, ["필드", "형식", "구분", "설명"], gateway_fields, [2450, 1300, 1350, 4260], font_size=8.5)

    add_page_break(doc)
    add_heading(doc, "9. 기준국·게이트웨이 결과 전송 예시", 1)
    add_code(doc, {
        "context": {
            "eventId": "10000000-0000-4000-8000-000000000001",
            "requestId": "c2a2fa8b-bc6c-480a-b92f-7020691a3f1e",
            "sourceSystem": "jininfra-lpwa-controller",
            "occurredAt": "2026-08-03T04:30:00Z",
            "schemaVersion": "1.0",
            "reportedByAssetId": "30000000-0000-4000-8000-000000000001",
            "reportingRole": "GATEWAY"
        },
        "data": {
            "assetId": "30000000-0000-4000-8000-000000000001",
            "observedAt": "2026-08-03T04:30:00Z",
            "operationalStatus": "ONLINE",
            "basePosition": {"type": "Point", "coordinates": [128.6889, 36.3503, 305.2]},
            "rtcmFormat": "RTCM3",
            "rtcmAvailable": True,
            "correctionAgeSeconds": 0.8,
            "deliveryMode": "BROADCAST",
            "beaconChannel": 1,
            "uplinkChannelCount": 7,
            "connectedTerminals": 4,
            "allocatedSlots": [
                {"terminalAssetId": "30000000-0000-4000-8000-000000000021", "channel": 2, "slot": 1}
            ],
            "ethernetBackhaul": {"connected": True, "type": "LTE", "latencyMs": 42},
            "attributes": {"controllerStatus": "NORMAL"}
        }
    })
    add_callout(doc, "제어 명령 처리", "기준국 셋업, 채널·슬롯 할당, FOTA 등의 명령 API는 진인프라가 실제 지원 명령·파라미터·ACK·오류코드를 회신한 뒤 v1.0 부록으로 확정합니다. 본 v0.9에서는 결과 수신 API를 우선 검토합니다.")
    add_delivery_rules(doc)
    add_vendor_decision_pages(
        doc,
        "주식회사 진인프라",
        "LPWA RF 게이트웨이 / 네트워크 제어기 / 인터페이스 SW / 별도 연동 프로그램 중 선택",
        "RTK 단말 위치·측위품질, 이동형 기준국의 RTCM 제공상태, LPWA/LTE 활성망·전환정보, "
        "게이트웨이 접속 단말·채널·슬롯·백홀 상태를 상위 시스템에 제공하는 방식을 검토합니다.",
    )
    path = OUT / "투비유니콘_진인프라_RTK_LPWA_상위연계_Northbound_인터페이스_협의안_v0.9.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    paths = [build_ndps(), build_jininfra()]
    for path in paths:
        print(path)
